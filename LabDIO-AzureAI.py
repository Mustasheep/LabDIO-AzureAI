pip install requests python-docx

import requests
from docx import Document
import os

subscription_key = "YOUR_SUB_KEY"
endpoint = 'https://link-do-endpoint.com'
location = "local"
target_language = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Sub-Key': subscription_key,
      'Ocp-Apim-Sub-Region': location,
      'X-ClientTraceId': str(os.urandom(16))
  }

  body = [{
      'text': text
  }]
  params = {
      'api-version': 'VERSÃO',
      'from': 'en',
      'to': target_language
  }
  request = request.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]
  
# Traduzindo documento
def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, target_language)
    full_text.append(translated_text)
  
  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", f"_{target_language}.docx")
  translated_doc.save(path_translated)
  return path_translated

# Inserindo Arquivo
input_file = "/content/NOMEARQUIVO.docx"
translate_document(input_file)